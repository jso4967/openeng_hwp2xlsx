from docx import Document
from docx.shared import RGBColor
import classes, model_core, re
from openpyxl import Workbook

DEGREE = 3
additional_msg1 = "\n\n=================================================================\n" \
                 "예외 ! <규칙에 맞지 않는 문자열 형태> : "
additional_msg2 = "\n=================================================================\n" \
                  "아래 사항들을 고려해보세요\n\n" \
                  "1. 파란색 줄과 검은색 줄이 확실하게 구분되었는가\n" \
                  "2. 해당 줄이 전부 검정색인가\n" \
                  "3. 선지번호가 한국어, 영어 동일하게 동일 개수로 인식되었는가\n" \
                  "4. 안된다면 word의 순서 및 형식 자동 완성 기능을 제거해보자(물결표시)\n" \
                  "5. 문장의 맨 첫글자가 숫자나 알파벳, 한국어 혹은 원문자가 아닌 경우(특수문자 인식불가)\n" \
                  "6. 문제번호 및 선지번호가 1번부터 순서대로 오름차순인가\n\n\n"

def is_blue_text(paragraph):
    if paragraph.text.find('\n') >= 0:
        print(model_core.errorMsg("", additional_msg1 + paragraph.text + " \n1번 시도 요망" + additional_msg2))
        raise model_core.RuleErrorEx

    for run in paragraph.runs:
        if run.text == "":
            continue
        elif run.font.color.rgb == RGBColor(0x00, 0x00, 0xFF) or run.font.color.rgb == RGBColor(0x00, 0x59, 0xFF):
            return True
    return False

def extract_problem_set(docx_path):
    '''
    일반 파일로부터 prolem_set의 형식에 맞게 데이터 추출하는 메소드

    :param docx_path: 일반 파일의 docx 버전
    :return: 일반 파일의 모든 것을 problem_set으로 반환
    '''

    document = Document(docx_path)
    problem_set = []
    current_problem_number = 0
    current_option_number = 0
    current_eng_sentence = ""
    current_kor_sentence = ""
    current_gender = None



    for index in range(0, int((document.tables[0]._cells.__len__() - 1)/DEGREE)):

        # 전 셀의 내용과 동일하다면 건너뛰기
        if document.tables[0].cell(0, (index - 1) * DEGREE).text.__len__() == document.tables[0].cell(0,
                                                                                                 index * DEGREE).text.__len__():
            continue
        lines = document.tables[0].cell(0, index * DEGREE).text

        # 첫 문장을 조건으로 쓸데 없는 셀들 건너뛰기
        if lines.find("문제") == 0:  # 문제로 시작한다면 다음 첫 문장 살피기
            print("문제 및 보기")
            continue
        elif lines.find("정답 해결") == 0:  # 정답 해결)로 시작한다면 다음 첫 문장 살피기
            print("정답 해석")
            continue
        elif lines.find("영치법 영한치환") >= 0:
            print("영치법 영한치환")
            continue
        elif lines.find("영치법 한영치환") >= 0:
            print("영치법 한영치환")
            continue
        elif lines.strip() == "" or lines.strip()[0] == '*' or model_core.is_alphabet(lines):  # 영어 단어 혹은 *으로 시작하는 부분
            continue

        paragraphs = document.tables[0].cell(0, index * DEGREE).paragraphs

        if lines.find("영어 원문 분석") >= 0:
            print("영어 원문 분석")
            current_problem_number = lines[:lines.find("영어 원문 분석")].strip()

            problem_set.append(classes.Problem(current_problem_number))
            print(current_problem_number)

            for entry in paragraphs:

                # 파란 글씨(설명) 혹은 빈 칸이면 건너뛰기
                if entry.text.strip() == "" or is_blue_text(entry):
                    continue

                if entry.text.find("영어 원문 분석") >= 0:  # 문제 번호가 맨 앞에 오는 경우
                    continue

                elif entry.text.find("빈칸에") >= 0:
                    continue

                # 원숫자 찾기
                message, current_option_number = model_core.does_start_with_circled_number(entry.text,
                                                                                                current_option_number)

                if message.find("첫 문자가 선지 번호인 경우") >= 0:  # 원숫자가 맞다면
                    print(message)

                    if current_option_number > 0:  # 이전 선지번호의 내용을 업데이트 하고, 현재 영문과 성별을 지운다.

                        try:
                            problem_set[problem_set.__len__() - 1].addsentence(current_eng_sentence.strip(),
                                                                               current_kor_sentence.strip(),
                                                                               current_gender)
                            current_eng_sentence = ""
                            current_gender = None

                        except model_core.RuleErrorEx:
                            print(model_core.errorMsg("",
                                                      additional_msg1 + entry.text +
                                                      "\n문제 번호에 대한 에러일 가능성이 큽니다."
                                                      + additional_msg2))

                    # 성별 찾기
                    line, current_gender = model_core.find_gender(entry.text)

                    current_eng_sentence += line
                    print("현재 문장 : " + current_eng_sentence)
                    continue

                elif message.find("선지번호가 존재하지 않는 경우") >= 0:
                    if model_core.is_alphabet(entry.text.strip()[0]) or entry.text.strip()[0].isdigit():  # 영어인 경우
                        current_eng_sentence += (" " + entry.text.strip())

                    else:
                        print(model_core.errorMsg("", additional_msg1 + entry.text + "\n 첫 문자가 특수문자일 가능성이 큽니다." + additional_msg2))
                        raise model_core.RuleErrorEx

                    print(message, entry.text)
                    print("현재 문장 : " + current_eng_sentence)

                else:
                    print(message, entry.text)
                    continue

            #  맨 마지막 선지 번호를 업데이트
            if current_eng_sentence:
                try:
                    problem_set[problem_set.__len__() - 1].addsentence(current_eng_sentence.strip(),
                                                                       current_kor_sentence.strip(),
                                                                       current_gender)
                    current_eng_sentence = ""
                    current_gender = None
                    current_option_number = 0

                except model_core.RuleErrorEx:
                    print(model_core.errorMsg("",
                                              additional_msg1 + entry.text +
                                              "\n문제 번호에 대한 에러일 가능성이 큽니다."
                                              + additional_msg2))

                # problem_set[problem_set.__len__() - 1].addsentence(current_eng_sentence.strip(), current_kor_sentence.strip(),
                #                                                    current_gender)
                # current_eng_sentence = ""
                # current_gender = None
                # current_option_number = 0

        elif lines.find("영문 해석") >= 0:
            print("영문 해석")

            # 문제 번호 찾기
            if lines[:lines.find("영문 해석")].strip() == current_problem_number:  # 영어 원문 분석과 문제번호가 짝인가?
                print(current_problem_number)
            else:
                print(model_core.errorMsg("", additional_msg1 + "\n" + lines + "\n 문제 번호 이상 가능성 농후" + additional_msg2))
                raise model_core.RuleErrorEx

            for entry in paragraphs:

                # 빈 칸이면 건너뛰기
                if entry.text.strip() == "" or entry.text.find("영문 해석") >= 0 or entry.text.find("빈칸에") >= 0 :
                    continue

                # 원숫자 찾기
                message, current_option_number = model_core.does_start_with_circled_number(entry.text,
                                                                                                current_option_number)

                if message.find("첫 문자가 선지 번호인 경우") >= 0:  # 원숫자가 맞다면
                    print(message)
                    if current_option_number > 0 :  # 이전 선지번호의 내용을 업데이트 하고, 현재 영문과 성별을 지운다.
                        sentences = problem_set[problem_set.__len__() - 1].getsentences()

                        try:

                            current_sentence = sentences[current_option_number - 1]

                            current_sentence.edit_sentence(current_eng_sentence, current_kor_sentence)

                            current_kor_sentence = ""
                            current_gender = None

                        except model_core.RuleErrorEx:

                            print(model_core.errorMsg("",
                                                      additional_msg1 + entry.text +
                                                      "\n선지번호에 대한 에러일 가능성이 큽니다."
                                                      + additional_msg2))

                    text = entry.text
                    # 한국어 성별이 존재하는 경우
                    if text[text.find("남"):].find(":") != -1:
                        sentence_text = text[text.find(":") + 1:]

                    elif text[text.find("여"):].find(":") != -1:
                        sentence_text = text[text.find(":") + 1:]
                    else:
                        sentence_text = text[2:]

                    current_kor_sentence += sentence_text.strip()
                    print("현재 문장 : " + current_kor_sentence)
                    continue

                elif message.find("선지번호가 존재하지 않는 경우") >= 0:
                    if len(re.findall(u'[\u3130-\u318F\uAC00-\uD7A3]+', entry.text)) or entry.text.strip()[0].isdigit():  # 한국어 혹은 숫자인 경우
                        current_kor_sentence += (" " + entry.text.strip())

                    else:
                        print(model_core.errorMsg("",
                                                  additional_msg1 + entry.text + "\n첫 문자가 특수문자일 가능성이 큽니다." + additional_msg2))
                        raise model_core.RuleErrorEx

                    print(message, entry.text)
                    print("현재 문장 : " + current_kor_sentence)

                else:
                    print(message, entry.text)
                    continue

            #  맨 마지막 선지 번호를 업데이트

            try:
                sentences = problem_set[problem_set.__len__() - 1].getsentences()
                current_sentence = sentences[sentences.__len__() - 1]
                current_sentence.edit_sentence(current_eng_sentence, current_kor_sentence)

                current_kor_sentence = ""
                current_option_number = 0

            except model_core.RuleErrorEx:

                print(model_core.errorMsg("",
                                          additional_msg1 + entry.text +
                                          "\n선지번호에 대한 에러일 가능성이 큽니다."
                                          + additional_msg2))

            if current_problem_number.find("~") >= 0:  # 숫자에 ~가 존재한다면
                current_problem_number = current_problem_number.strip()
                start = int(current_problem_number[:current_problem_number.find("~")])
                end = int(current_problem_number[current_problem_number.find("~")+1:])

                for i in range(0, end - start):
                    problem_set.append(classes.Problem(current_problem_number))

                    for sentence in problem_set[problem_set.__len__() - 2].getsentences():
                        problem_set[problem_set.__len__() - 1].addsentence(sentence.get_sentence()[0], sentence.get_sentence()[1], sentence.get_gender())


        else :  # 규칙에 맞지 않는 상황
            print(model_core.errorMsg("", additional_msg1 + "\n" + lines + "\n" + additional_msg2))
            raise model_core.RuleErrorEx
            continue

    return problem_set

def save_as_sentence_xlsx(problem_for_sentence, save_path):

    wb = Workbook()
    dest_filename = save_path[:save_path.rfind(".")] + "_Sentence_result.xlsx"

    ws1 = wb.active
    ws1.title = "영문해석"
    row = 1
    flag = 0
    for problem in problem_for_sentence:

        problem_number = problem.get_problem_num()
        if repr(flag) == problem_number:
            continue
        flag = problem_number

        sentences = problem.getsentences()

        for idx, sentence in enumerate(sentences):
            ws1['A'+repr(row)] = row
            ws1['C'+repr(row)] = problem_number
            ws1['D'+repr(row)] = idx + 1

            if sentence.get_gender():
                ws1['E'+repr(row)] = sentence.get_gender()

            ws1['F'+repr(row)] = sentence.get_sentence()[0].strip()
            ws1['G'+repr(row)] = sentence.get_sentence()[1].strip()
            row += 1

    wb.save(filename = dest_filename)

def save_as_phrase_xlsx(problem_for_phrase, save_path):

    wb = Workbook()
    dest_filename = save_path[:save_path.rfind(".")] + "_Phrase_result.xlsx"

    ws1 = wb.active
    ws1.title = "영치법문장"

    row = 1
    for problem in problem_for_phrase:
        sentences = problem.getsentences()

        for sentence in sentences:
            for phrase in sentence.getphrases():
                temp_row = row

                for eng_phrase in phrase.get_eng_phrase():
                    ws1['D' + repr(temp_row)] = eng_phrase.strip()
                    temp_row += 1

                temp_row = row
                for kor_phrase in phrase.get_kor_phrase():
                    ws1['E' + repr(temp_row)] = kor_phrase.strip()
                    temp_row += 1
                row = temp_row

    wb.save(filename = dest_filename)






