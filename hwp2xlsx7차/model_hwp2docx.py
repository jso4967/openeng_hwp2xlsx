from docx import Document
import classes, re, model_core

def extract_phrase_set(docx_path):
    '''
    파일에서 phrase만 추출하는 메소드

    :param docx_path: docx-file to extract
    :return: only phrase values of file
    '''
    document = Document(docx_path)
    lines = document.paragraphs

    current_gender = None
    current_eng_sentence = ""
    current_kor_sentence = ""
    current_eng_phrases = []
    current_kor_phrases = []
    current_option_number = 0
    current_problem_set = []
    current_problem_number = 1

    for line in lines:
        text = line.text
        text = text.strip()

        if text == "":
            continue

        # 첫 문자가 문제 번호인 경우
        if (text[:text.find('.')]).isdecimal():
            print("첫 문자가 문제 번호인 경우 : ", text)
            current_gender = None
            current_eng_phrases = []
            current_kor_phrases = []
            current_problem_number = int(text[:text.find('.')]) - 1
            current_option_number = 0
            current_problem_set.append(
                classes.Problem(current_problem_number))  # 문제번호가 적힌 문장에서 항상 첫번째 글자는 숫자( = 문제번호) 여야 한다.

            continue

        # 첫 문자가 선지 번호인 경우
        message, current_option_number = model_core.does_start_with_circled_number(text, current_option_number)

        if message.find("첫 문자가 선지 번호인 경우") >= 0 :  # 문장이 올바른 선지 번호로 시작하는 경우
            print(message)

            # 성별 찾기
            sentence_text, current_gender = model_core.find_gender(text)

            sentence_text = sentence_text.strip()
            current_eng_sentence = sentence_text
            current_eng_phrases = sentence_text.split('/')

            #  각 구에 대하여 공백을 지운다.
            current_eng_phrases = model_core.delete_empty_space(current_eng_phrases)
            print(current_eng_phrases)

            continue

        # 첫 문자가 문제 구분자인 경우 혹은 문장에 문자가 없는 경우
        if text.find("===") >= 0 or (text == ''):
            print("첫 문자가 문제 구분자 이거나 아무 문자도 없을 때 : ", text)
            continue

        # 첫 문자가 한글인 경우

        if len(re.findall(u'[\u3130-\u318F\uAC00-\uD7A3]+', text)):
            print("첫 문자가 한글인 경우 : ", text)
            current_kor_sentence = text
            current_kor_phrases = current_kor_sentence.split('/')


            #  각 구에 대하여 공백을 지운다.
            current_kor_phrases = model_core.delete_empty_space(current_kor_phrases)
            print(current_kor_phrases)

            #  해당 문제에 문장 추가
            try:
                last_option_number = current_problem_set[current_problem_number].getsentences().__len__() - 1
                if current_option_number == last_option_number:
                    try:
                        current_problem_set[current_problem_number].getsentences()[last_option_number].edit_sentence(
                            current_eng_sentence, current_kor_sentence)
                    except model_core.RuleErrorEx:
                        print(model_core.errorMsg(text + " \n선지 번호에 관한 에러일 가능성이 큽니다."))
                else:
                    current_problem_set[current_problem_number].addsentence(current_eng_sentence, current_kor_sentence,
                                                                            current_gender)
            except model_core.RuleErrorEx:
                print(model_core.errorMsg(text + " \n문제 번호에 관한 에러일 가능성이 큽니다."))

            if current_eng_phrases.__len__() == current_kor_phrases.__len__():
                print("++++++++++++++++++++++++++동일++++++++++++++++++++++++++")
                #  해당 문장에 구 추가
                try:
                    if last_option_number == current_option_number:  # 기존 구에 덧붙이기
                        (current_problem_set[current_problem_number].getsentences())[current_option_number].edit_phrase(
                            current_eng_phrases, current_kor_phrases)
                    else:  # 새로 구 만들기
                        (current_problem_set[current_problem_number].getsentences())[current_option_number].addphrase(
                            current_eng_phrases, current_kor_phrases, None)
                except model_core.RuleErrorEx:
                    print(model_core.errorMsg(text + " \n문제 혹은 선지번호에 관한 에러일 가능성이 큽니다."))
            else:
                print(model_core.errorMsg(text + " \n슬래쉬 갯수가 안맞는 경우"))
                raise model_core.RuleErrorEx
            try:
                if current_problem_set[current_problem_number].getsentences():
                    print("현재 문제번호 : " + str(current_problem_number + 1) + " 현재 선지번호 : " + str(current_option_number + 1))
                    try:
                        for entry in (current_problem_set[current_problem_number].getsentences())[
                            current_option_number].getphrases():
                            print(entry.getall())
                    except model_core.RuleErrorEx:
                        print(model_core.errorMsg(text + " \n선지 번호에 관한 에러일 가능성이 큽니다."))
            except model_core.RuleErrorEx:
                print(model_core.errorMsg(text + " \n문제 번호에 관한 에러일 가능성이 큽니다."))
            continue

        # 첫 문자가 영어이거나 숫자인 경우
        if text.strip() == "" or text.strip()[0].isdigit() or model_core.is_alphabet(text) :
            current_eng_sentence = text.strip()
            current_eng_phrases = current_eng_sentence.split('/')

            #  각 구에 대하여 공백을 지운다.
            current_eng_phrases = model_core.delete_empty_space(current_eng_phrases)
            print(current_eng_phrases)

            continue

        # 위에 해당하지 않는 경우
        print(model_core.errorMsg(text+"\n 일반적으로 첫글자가 특수문자인 경우"))
        raise model_core.RuleErrorEx

    return current_problem_set

def construct_phrase_data(problem_set):
    '''

    :param problem_set: 여러개의 문제형식으로 되어있는 입력값
    :return: 저장할 파일의 데이터를 작성
    '''
    data_eng_kor = ""
    data_kor_eng = ""
    problem_number = 1
    for problem in problem_set:

        # 문제 번호 출력
        data_eng_kor += "\n" + str(problem_number) + "\n"
        data_kor_eng += "\n" + str(problem_number) + "\n"

        problem_number += 1
        option_number = 0xe291a0
        flag = 1
        for sentence in problem.getsentences():

            # 문항 번호 출력
            data_eng_kor += bytes.fromhex(str(hex(option_number)[2:])).decode('utf-8') + " "
            data_kor_eng += bytes.fromhex(str(hex(option_number)[2:])).decode('utf-8') + " "
            option_number += 1

            if sentence.getphrases():
                phrase = sentence.getphrases()[0]
                eng_phrase = phrase.get_eng_phrase()
                kor_phrase = phrase.get_kor_phrase()
                if len(eng_phrase) == len(kor_phrase):
                    index = len(eng_phrase)
                    substituted_sentence_eng_kor = ""
                    substituted_sentence_kor_eng = ""

                    for index in range(index):
                        if eng_phrase[index] == "":
                            continue
                        if flag == 1:
                            substituted_sentence_eng_kor += (eng_phrase[index] + " ")
                            substituted_sentence_kor_eng += (kor_phrase[index] + " ")
                            flag = 0
                        else:
                            substituted_sentence_eng_kor += (kor_phrase[index] + " ")
                            substituted_sentence_kor_eng += (eng_phrase[index] + " ")
                            flag = 1

                    data_eng_kor += substituted_sentence_eng_kor + "\n"
                    data_kor_eng += substituted_sentence_kor_eng + "\n"
            flag = (option_number + 1) % 2

        option_number = 0xe291a0

    return (data_eng_kor, data_kor_eng)

def save_phrase_file(constructed_data, input_path, output_path):
    '''
    작성된 파일의 데이터를 실제 파일에 작성하여 저장하는 함수
    :param constructed_data: 파일에 쓸 값
    :param input_path: 입력 파일의 경로로부터 파일의 이름을 가져온다.
    :param output_path: 저장 경로
    :return:
    '''
    # 파일 세팅

    ## 파일 저장 경로 설정
    path1 = output_path + input_path[input_path.rfind('/'):input_path.find('.')] + "_영한치환버전.docx"
    path2 = output_path + input_path[input_path.rfind('/'):input_path.find('.')] + "_한영치환버전.docx"

    # docx 파일로 만들기

    document = Document()

    document.add_paragraph(constructed_data[0])
    document.save(path1)

    document = Document()

    document.add_paragraph(constructed_data[1])
    document.save(path2)

# def check_number_of_slashes(text1, text2):
#     '''
#     두 문장의 슬래쉬 갯수를 확인하여 에러를 더욱 구체화한다.
#     :param text1: 비교할 첫 번째 문자열
#     :param text2: 비교할 두 번째 문자열
#     :return: Bool 자료형으로 True or False
#     '''
#     if text1.split('/') != text2.split('/')