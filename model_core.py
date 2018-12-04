from docx import Document
import classes
import re

# TODO : phrase 빈칸 삭제!

class RuleErrorEx(BaseException):pass

def errorMsg(text):
    msg = "=============================================================\n"
    msg += "예외 규칙에 맞지 않는 문자 형태! : " + text
    msg += "\n===========================================================\n"
    msg += "1. 선지번호가 등장하는 문장은 항상 맨 처음 선지번호, 띄어쓰기, 이후 문장 순으로 적혀있어야 한다.\n"
    msg += "2. 문제번호가 적힌 문장에서 항상 첫번째 글자는 숫자( = 문제번호) 여야 한다.\n"
    msg += "3. 특수문자가 맨 앞에 있는 지문인 경우 1번 선지 번호 밑에 붙여넣는다."
    msg += "4. 특수문자가 맨 앞에 있는 이어지는 문장인 경우 항상 영어 혹은 한글 혹은 숫자로만 시작되어야 한다."

    return msg

def convert_for_kor2eng(path = None):
    document = Document(path)
    lines = document.paragraphs

    current_gender = None
    current_eng_sentence = ""
    current_kor_sentence = ""
    current_eng_phrase = []
    current_kor_phrase = []
    current_option_number = 0
    current_problem_set = []
    current_problem_number = 1

    for line in lines:
        text = line.text
        text = text.strip()
        # 첫 문자가 문제 번호인 경우
        if (text[:text.find('.')]).isdecimal():
            print("첫 문자가 문제 번호인 경우 : ", text)
            current_gender = None
            current_eng_phrase = []
            current_kor_phrase = []
            current_problem_number = int(text[:text.find('.')]) - 1
            current_option_number = 0
            current_problem_set.append(
                classes.Problem(current_problem_number))  # 문제번호가 적힌 문장에서 항상 첫번째 글자는 숫자( = 문제번호) 여야 한다.

            continue

        # 첫 문자가 선지 번호인 경우
        option_number = 0xe291a0
        message = ""
        for i in range(50):
            # 헥스코드를 이용하여 선지번호인지 체크

            # 20인 부분에서 점프 : 유니코드 때문에.. https://unicode-table.com/en/2473/
            if option_number == 14848436:
                option_number += 63453

            # 35인 부분에서 점프 : 유니코드 때문에.. https://unicode-table.com/en/32B1/
            if option_number == 14911904:
                option_number += 273

            is_option_number = text[:2].find(bytes.fromhex(str(hex(option_number)[2:])).decode('utf-8'))
            if is_option_number >= 0 and (current_option_number == 0 or current_option_number + 1 == i):
                message = "첫 문자가 선지 번호인 경우 : " + text
                current_option_number = i
                break
            elif is_option_number >= 0 and current_option_number + 1 != i:
                break

            option_number += 1

        if message:
            print(message)
            textforsplit = ""
            # 성별 찾기
            textforgender = text.strip()
            if textforgender[textforgender.find("M"):].find(":") != -1:
                current_gender = "M"
                textforsplit = text[textforgender.find(":") + 1:]
                print(textforsplit + " gender is " + current_gender)
            elif textforgender[textforgender.find("W"):].find(":") != -1:
                current_gender = "W"
                textforsplit = text[textforgender.find(":") + 1:]
                print(textforsplit + " gender is " + current_gender)
            else:
                textforsplit = text[2:]  # 선지번호가 등장하는 문장은 항상 맨 처음 선지번호, 띄어쓰기, 이후 문장 순으로 적혀있어야 한다.

            textforsplit = textforsplit.strip()
            current_eng_sentence = textforsplit
            current_eng_phrase = textforsplit.split('/')

            #  각 구에 대하여 공백을 지운다.
            i = current_eng_phrase.__len__() - 1
            while i >= 0:
                current_eng_phrase[i] = current_eng_phrase[i].strip()
                if current_eng_phrase[i] == "":
                    del current_eng_phrase[i]
                i -= 1

            print(current_eng_phrase)
            continue

        # 첫 문자가 문제 구분자인 경우 혹은 문장에 문자가 없는 경우
        if text.find("===") >= 0 or (text == ''):
            print("첫 문자가 문제 구분자 이거나 아무 문자도 없을 때 : ", text)
            continue

        # 첫 문자가 한글인 경우

        if len(re.findall(u'[\u3130-\u318F\uAC00-\uD7A3]+', text)):
            print("첫 문자가 한글인 경우 : ", text)
            current_kor_sentence = text.strip()
            current_kor_phrase = current_kor_sentence.split('/')

            #  각 구에 대하여 공백을 지운다.
            i = current_kor_phrase.__len__() - 1
            while i >= 0:
                current_kor_phrase[i] = current_kor_phrase[i].strip()
                if current_kor_phrase[i] == "":
                    del current_kor_phrase[i]
                i -= 1

            print(current_kor_phrase)
            #  해당 문제에 문장 추가
            last_option_number = current_problem_set[current_problem_number].getsentences().__len__() - 1
            if current_option_number == last_option_number:
                current_problem_set[current_problem_number].getsentences()[last_option_number].edit_sentence(
                    current_eng_sentence, current_kor_sentence)
            else:
                current_problem_set[current_problem_number].addsentence(current_eng_sentence, current_kor_sentence,
                                                                        current_gender)

            if current_eng_phrase.__len__() == current_kor_phrase.__len__():
                print("++++++++++++++++++++++++++동일++++++++++++++++++++++++++")
                #  해당 문장에 구 추가

                if last_option_number == current_option_number:  # 기존 구에 덧붙이기
                    (current_problem_set[current_problem_number].getsentences())[current_option_number].edit_phrase(current_eng_phrase, current_kor_phrase)
                else:  # 새로 구 만들기
                    (current_problem_set[current_problem_number].getsentences())[current_option_number].addphrase(
                        current_eng_phrase, current_kor_phrase, None)
            # else:
            #     (current_problem_set[current_problem_number].getsentences())[current_option_number].addphrase(current_eng_phrase, current_kor_phrase, None)
            #     pass

            if current_problem_set[current_problem_number].getsentences():
                print("현재 문제번호 : " + str(current_problem_number + 1) + " 현재 선지번호 : " + str(current_option_number + 1))
                for entry in (current_problem_set[current_problem_number].getsentences())[
                    current_option_number].getphrases():
                    print(entry.getall())
            continue

        # 첫 문자가 영어이거나 숫자인 경우
        isAlphabet = ord(text.strip()[0])
        if (65 <= isAlphabet and isAlphabet <= 90) or (97 <= isAlphabet and isAlphabet <= 122) or text[0].isdigit():
            print("첫 문자가 영어인 경우")
            print(text)
            current_eng_sentence = text.strip()
            current_eng_phrase = current_eng_sentence.split('/')

            #  각 구에 대하여 공백을 지운다.
            i = current_eng_phrase.__len__() - 1
            while i >= 0:
                current_eng_phrase[i] = current_eng_phrase[i].strip()

                if current_eng_phrase[i] == "":
                    del current_eng_phrase[i]

                i -= 1

            print(current_eng_phrase)
            continue

        # 위에 해당하지 않는 경우
        # TODO : 규칙 정리해야 함
        print(errorMsg(text))
        raise RuleErrorEx
        continue

    return current_problem_set

def make_string(problem_set):
    data_eng_kor = ""
    data_kor_eng = ""
    problem_number = 1
    for problem in problem_set:

        # 문제 번호 출력
        data_eng_kor += "\n" + str(problem_number) + "\n"
        data_kor_eng += "\n" + str(problem_number) + "\n"

        problem_number += 1
        option_number = 0xe291a0
        flag = 1
        for sentence in problem.getsentences():

            # 문항 번호 출력
            data_eng_kor += bytes.fromhex(str(hex(option_number)[2:])).decode('utf-8') + " "
            data_kor_eng += bytes.fromhex(str(hex(option_number)[2:])).decode('utf-8') + " "
            option_number += 1

            if sentence.getphrases():
                phrase = sentence.getphrases()[0]
                eng_phrase = phrase.get_eng_phrase()
                kor_phrase = phrase.get_kor_phrase()
                if len(eng_phrase) == len(kor_phrase):
                    index = len(eng_phrase)
                    substituted_sentence_eng_kor = ""
                    substituted_sentence_kor_eng = ""

                    for index in range(index):
                        if eng_phrase[index] == "":
                            continue
                        if flag == 1:
                            substituted_sentence_eng_kor += (eng_phrase[index] + " ")
                            substituted_sentence_kor_eng += (kor_phrase[index] + " ")
                            flag = 0
                        else:
                            substituted_sentence_eng_kor += (kor_phrase[index] + " ")
                            substituted_sentence_kor_eng += (eng_phrase[index] + " ")
                            flag = 1

                    data_eng_kor += substituted_sentence_eng_kor + "\n"
                    data_kor_eng += substituted_sentence_kor_eng + "\n"
            flag = (option_number + 1) % 2

        option_number = 0xe291a0

    return data_eng_kor, data_kor_eng

def make_final_file(data1, data2, input_path, output_path):
    # 파일 세팅
    path1 = output_path  + input_path[input_path.rfind('/'):input_path.find('.')] + "_영한치환버전.docx"
    path2 = output_path  + input_path[input_path.rfind('/'):input_path.find('.')] + "_한영치환버전.docx"

    # docx 파일로 만들기

    document = Document()

    document.add_paragraph(data1)
    document.save(path1)

    document = Document()

    document.add_paragraph(data2)
    document.save(path2)